"""Build the frozen Los Angeles surface-heat research paper.

This script is publication-only. It reads already-published development and
2025 final-evaluation summaries and figures. It does not fit a model, rebuild a
target, open a remote asset, or alter any canonical evaluation artifact.

Design contract
---------------
Preset: ``narrative_proposal``
Header pattern: ``editorial_cover``
Page: US Letter portrait, 1-inch margins, 0.492-inch header/footer distances,
      6.5-inch / 9360-DXA content width.
Body: Calibri 11 pt, justified, 0 pt before, 8 pt after, 1.333 spacing.
H1: 16 pt #2E74B5, 18 pt before, 10 pt after.
H2: 13 pt #2E74B5, 12 pt before, 6 pt after.
H3: 12 pt #1F4D78, 8 pt before, 4 pt after.
Lists: 0.181-inch marker position, 0.375-inch text indent,
       0.194-inch hanging indent, 4 pt after, 1.208 spacing.
Tables: 9360-DXA fixed width, 120-DXA indent, 80/80/120/120-DXA
        top/bottom/start/end cell margins, #F4F6F9 header fill.

Named visual overrides
----------------------
``editorial_cover_accent`` uses dark ink #18201F, terracotta #C9583E,
and forest #193C36 for the title page and restrained lead callouts.
``figure_caption`` uses 9 pt dark gray, left aligned, 4 pt before/8 pt after.
``source_note`` uses 8 pt muted gray, left aligned, 0 pt before/8 pt after.
"""

from __future__ import annotations

import csv
import json
import sys
from collections.abc import Iterable, Sequence
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image, ImageOps
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Spacer,
    TableStyle,
)
from reportlab.platypus import (
    Image as PdfImage,
)
from reportlab.platypus import (
    Paragraph as PdfParagraph,
)
from reportlab.platypus import (
    Table as PdfTable,
)

ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / "data" / "processed" / "final_test_2025" / "final_evaluation"
DEV_TABLES = ROOT / "reports" / "tables"
FIGURES = ROOT / "reports" / "figures" / "generated"
ASSET_DIR = ROOT / "tmp" / "publication_assets"
TMP_DIR = ROOT / "tmp" / "publication_paper"
OUTPUT_DIR = ROOT / "exports" / "PUBLICATION_MATERIALS"
OUTPUT_DOCX = OUTPUT_DIR / "LA_Surface_Heat_Research_Paper.docx"
OUTPUT_PDF = OUTPUT_DIR / "LA_Surface_Heat_Research_Paper.pdf"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x18, 0x20, 0x1F)
TERRACOTTA = RGBColor(0xC9, 0x58, 0x3E)
FOREST = RGBColor(0x19, 0x3C, 0x36)
MUTED = RGBColor(0x5D, 0x65, 0x63)
LIGHT_MUTED = RGBColor(0x86, 0x8D, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_FILL = "F4F6F9"
RULE = "D7DBE2"
CALLOUT_FILL = "F7EEE9"
FOREST_FILL = "EAF1EF"


def read_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_border(cell, *, color: str = RULE, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def apply_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    widths = [int(value) for value in widths_dxa]
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    _set_width(tbl_pr, "w:tblW", TABLE_WIDTH_DXA)
    indent = _ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = Twips(width)
    for row in table.rows:
        row.height = None
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for col_idx, cell in enumerate(row.cells):
            cell.width = Twips(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            _set_width(tc_pr, "w:tcW", widths[col_idx])
            margins = _ensure_child(tc_pr, "w:tcMar")
            for side, value in CELL_MARGINS_DXA.items():
                margin = _ensure_child(margins, f"w:{side}")
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(value))
            set_cell_border(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    repeat_table_header(row)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor = INK,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    size: float = 9.2,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(
    doc: Document,
    *,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_dxa: Sequence[int],
    numeric_columns: set[int] | None = None,
) -> object:
    numeric_columns = numeric_columns or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], TABLE_FILL)
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            color=FOREST,
            align=(
                WD_ALIGN_PARAGRAPH.CENTER
                if idx in numeric_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            ),
            size=9.1,
        )
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(
                cells[idx],
                str(value),
                align=(
                    WD_ALIGN_PARAGRAPH.CENTER
                    if idx in numeric_columns
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
            )
    apply_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.page_break_before = False

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = INK
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(9)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.italic = False
    subtitle.font.color.rgb = FOREST
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0x3B, 0x43, 0x41)
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.05
    caption.paragraph_format.keep_together = True

    if "Source Note" not in styles:
        styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    source_note = styles["Source Note"]
    source_note.font.name = "Calibri"
    source_note.font.size = Pt(8)
    source_note.font.color.rgb = MUTED
    source_note._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    source_note._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    source_note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source_note.paragraph_format.space_before = Pt(0)
    source_note.paragraph_format.space_after = Pt(4)
    source_note.paragraph_format.line_spacing = 1.0
    source_note.paragraph_format.keep_together = True

    if "Table Caption" not in styles:
        styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    table_caption = styles["Table Caption"]
    table_caption.font.name = "Calibri"
    table_caption.font.size = Pt(9.5)
    table_caption.font.bold = True
    table_caption.font.color.rgb = FOREST
    table_caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    table_caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    table_caption.paragraph_format.space_before = Pt(4)
    table_caption.paragraph_format.space_after = Pt(4)
    table_caption.paragraph_format.keep_with_next = True
    table_caption.paragraph_format.keep_together = True

    if "Reference" not in styles:
        styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    reference = styles["Reference"]
    reference.font.name = "Calibri"
    reference.font.size = Pt(9.2)
    reference.font.color.rgb = INK
    reference._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    reference._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    reference.paragraph_format.left_indent = Inches(0.30)
    reference.paragraph_format.first_line_indent = Inches(-0.30)
    reference.paragraph_format.space_after = Pt(4)
    reference.paragraph_format.line_spacing = 1.08


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def configure_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("LA NEIGHBORHOOD SURFACE HEAT  /  RESEARCH PAPER")
    set_run_font(run, size=8.2, color=MUTED, bold=True)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)


def create_numbering(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "279")
    p_pr.append(indentation)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_body_paragraph(
    doc: Document,
    text: str,
    *,
    bold_lead: str | None = None,
    keep_together: bool = False,
) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_together = keep_together
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)
    return paragraph


def add_callout(doc: Document, label: str, text: str, *, forest: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.17)
    paragraph.paragraph_format.right_indent = Inches(0.17)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.18
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), FOREST_FILL if forest else CALLOUT_FILL)
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), "193C36" if forest else "C9583E")
    border.append(left)
    p_pr.append(border)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(
        label_run,
        size=10.4,
        color=FOREST if forest else TERRACOTTA,
        bold=True,
    )
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.4, color=INK)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Figure Caption")
    paragraph.add_run(text)


def add_table_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Table Caption")
    paragraph.add_run(text)


def add_source_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Source Note")
    paragraph.add_run(text)


def add_figure(
    doc: Document,
    path: Path,
    *,
    width_inches: float,
    alt_text: str,
    caption: str,
    source_note: str,
) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = paragraph._p.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", alt_text)
        doc_pr[0].set("title", alt_text[:100])
    add_caption(doc, caption)
    add_source_note(doc, source_note)


def build_map_composite() -> Path:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    observed = ASSET_DIR / "final-map-hires-1.png"
    m2 = ASSET_DIR / "final-map-hires-3.png"
    b1_residual = ASSET_DIR / "final-map-4.png"
    m2_residual = ASSET_DIR / "final-map-5.png"
    paths = [observed, m2, b1_residual, m2_residual]
    for path in paths:
        if not path.exists():
            raise FileNotFoundError(path)
    images = [Image.open(path).convert("RGB") for path in paths]
    target = (900, 900)
    cells = [
        ImageOps.pad(image, target, method=Image.Resampling.LANCZOS, color="white")
        for image in images
    ]
    gutter = 18
    canvas = Image.new(
        "RGB",
        (target[0] * 2 + gutter, target[1] * 2 + gutter),
        "white",
    )
    canvas.paste(cells[0], (0, 0))
    canvas.paste(cells[1], (target[0] + gutter, 0))
    canvas.paste(cells[2], (0, target[1] + gutter))
    canvas.paste(cells[3], (target[0] + gutter, target[1] + gutter))
    output = TMP_DIR / "final_map_composite.png"
    canvas.save(output, format="PNG", optimize=True, dpi=(200, 200))
    return output


def load_metrics() -> dict:
    final_metrics = {row["model_id"]: row for row in read_csv(FINAL / "model_metrics.csv")}
    bootstrap = read_json(FINAL / "crossed_bootstrap.json")
    hotspot = {row["model_id"]: row for row in read_csv(FINAL / "hotspot_summary.csv")}
    sensors = read_csv(FINAL / "sensor_summary.csv")
    sentinel = read_csv(FINAL / "sentinel_stratum_summary.csv")
    gates = read_csv(FINAL / "protocol_gates.csv")
    per_date = read_csv(FINAL / "per_date_metrics.csv")
    development = read_json(
        DEV_TABLES / "model_results_initial" / "model_results_initial_summary.json"
    )
    development_endpoints = read_json(
        DEV_TABLES
        / "model_endpoint_diagnostics"
        / "model_endpoint_diagnostics_summary.json"
    )
    b1_by_date = {
        row["target_date"]: row for row in per_date if row["model_id"] == "B1"
    }
    m2_by_date = {
        row["target_date"]: row for row in per_date if row["model_id"] == "M2"
    }
    dates = sorted(set(b1_by_date) & set(m2_by_date))
    lower_abs = sum(
        float(m2_by_date[date]["mae_c"]) < float(b1_by_date[date]["mae_c"])
        for date in dates
    )
    lower_anomaly = sum(
        float(m2_by_date[date]["within_date_anomaly_mae_c"])
        < float(b1_by_date[date]["within_date_anomaly_mae_c"])
        for date in dates
    )
    higher_spearman = sum(
        float(m2_by_date[date]["spearman_rho"])
        > float(b1_by_date[date]["spearman_rho"])
        for date in dates
    )
    return {
        "final": final_metrics,
        "bootstrap": bootstrap,
        "hotspot": hotspot,
        "sensors": sensors,
        "sentinel": sentinel,
        "gates": gates,
        "per_date": per_date,
        "dates": dates,
        "lower_abs_dates": lower_abs,
        "lower_anomaly_dates": lower_anomaly,
        "higher_spearman_dates": higher_spearman,
        "development": development,
        "development_endpoints": development_endpoints,
    }


def add_cover(doc: Document, metrics: dict) -> None:
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(17)
    run = kicker.add_run("URBAN CLIMATE  /  HISTORICAL HINDCAST")
    set_run_font(run, size=10, color=TERRACOTTA, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run(
        "Can Public Weather, Land-Use, and Lagged Optical Satellite Features "
        "Predict Neighborhood-Scale Surface Heat?"
    )
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A Prespecified Historical Hindcast in Los Angeles")

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(24)
    run = rule.add_run("\u2501" * 22)
    set_run_font(run, size=8, color=TERRACOTTA)

    result = doc.add_paragraph()
    result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    result.paragraph_format.space_after = Pt(7)
    result.paragraph_format.keep_together = True
    big = result.add_run(
        f"{metrics['bootstrap']['relative_mae_improvement_percent']:.1f}%"
    )
    set_run_font(big, size=25, color=FOREST, bold=True)
    result.add_run("\n")
    small = result.add_run("lower held-out point-estimate MAE for M2 vs. B1")
    set_run_font(small, size=11, color=INK, bold=True)

    qualifier = doc.add_paragraph()
    qualifier.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qualifier.paragraph_format.space_before = Pt(10)
    qualifier.paragraph_format.space_after = Pt(46)
    qualifier.paragraph_format.keep_together = True
    q = qualifier.add_run(
        "Promising held-out predictive signal; not protocol-confirmed because "
        "the prespecified 95% interval crossed zero."
    )
    set_run_font(q, size=11.5, color=TERRACOTTA, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)
    r = meta.add_run("Research paper  |  Frozen final evaluation  |  28 July 2026")
    set_run_font(r, size=10, color=MUTED, bold=True)
    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_after = Pt(0)
    r = scope.add_run(
        "City of Los Angeles census tracts  |  Landsat overpass dates  |  2020-2025"
    )
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    metrics = load_metrics()
    map_composite = build_map_composite()
    b1 = metrics["final"]["B1"]
    m2 = metrics["final"]["M2"]
    boot = metrics["bootstrap"]
    dev = metrics["development"]["primary_comparison"]
    dev_gates = metrics["development"]["protocol_success_gates"]
    dev_hot = metrics["development_endpoints"]["relative_endpoint"][
        "focus_joint_models"
    ]

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_page(section)
        configure_header_footer(section)
    bullet_num_id = create_numbering(doc, kind="bullet")
    add_cover(doc, metrics)

    doc.add_heading("Abstract", level=1)
    abstract = (
        "Urban surface heat varies sharply within cities, but predicting "
        "neighborhood-scale patterns without target-day thermal observations "
        "remains challenging. We tested whether public weather, land-use and "
        "geographic, calendar, and lagged nonthermal optical features could "
        "predict QA-filtered daytime Landsat land-surface temperature (LST) "
        "across 1,096 Los Angeles census tracts. LST was treated as a clear-sky "
        "surface-heat hazard proxy, not air temperature or human exposure. This "
        "was a historical hindcast: predictions originated at 00:00 local time "
        "on each target date, and dynamic observed inputs ended at day -1. A "
        "ridge baseline using 23 calendar and lagged-weather features (B1) was "
        "compared with a histogram-gradient-boosting model using 46 features "
        "(M2). Development evaluation covered 63,403 tract-date observations, "
        "65 independent dates, and 71 spatial blocks from 2020-2024. M2 reduced "
        f"development MAE from {dev['baseline_point_mae_c']:.4f} C to "
        f"{dev['target_model_point_mae_c']:.4f} C, a "
        f"{dev['relative_mae_improvement_percent']:.2f}% improvement "
        f"(crossed date-by-block 95% CI: "
        f"{dev['relative_mae_improvement_ci_lower_percent']:.2f}% to "
        f"{dev['relative_mae_improvement_ci_upper_percent']:.2f}%). Models and "
        "thresholds were then frozen before a one-time 2025 test. The final "
        "cohort contained 15,116 observations across 15 usable dates and 71 "
        f"blocks. Held-out MAE was {float(b1['equal_date_weighted_mae_c']):.4f} "
        f"C for B1 and {float(m2['equal_date_weighted_mae_c']):.4f} C for M2, "
        f"a point reduction of {boot['absolute_mae_improvement_c']:.4f} C or "
        f"{boot['relative_mae_improvement_percent']:.2f}%. However, the "
        "prespecified 5,000-replicate crossed-cluster interval ranged from "
        f"{boot['relative_mae_improvement_ci_lower_percent']:.2f}% to "
        f"{boot['relative_mae_improvement_ci_upper_percent']:.2f}%, so the "
        "required positive lower-bound gate failed. M2 nevertheless showed "
        "higher neighborhood rank correlation and stronger top-20% hotspot "
        "identification. These findings provide promising held-out predictive "
        "evidence, but they do not establish protocol-level confirmation, "
        "causation, human heat exposure, or operational forecasting performance."
    )
    add_body_paragraph(doc, abstract)
    add_callout(
        doc,
        "Interpretation boundary",
        "The outcome is clear-sky daytime land-surface temperature. It is not "
        "near-surface air temperature, personal exposure, heat illness, or "
        "mortality.",
    )

    doc.add_heading("Key findings", level=2)
    add_list_item(
        doc,
        f"Development (2020-2024): M2 improved equal-date MAE by "
        f"{dev['relative_mae_improvement_percent']:.2f}% "
        f"(95% CI {dev['relative_mae_improvement_ci_lower_percent']:.2f}% to "
        f"{dev['relative_mae_improvement_ci_upper_percent']:.2f}%).",
        bullet_num_id,
    )
    add_list_item(
        doc,
        f"Frozen 2025 test: M2 reduced point-estimate MAE by "
        f"{boot['relative_mae_improvement_percent']:.2f}%, but the 95% interval "
        f"({boot['relative_mae_improvement_ci_lower_percent']:.2f}% to "
        f"{boot['relative_mae_improvement_ci_upper_percent']:.2f}%) crossed zero.",
        bullet_num_id,
    )
    add_list_item(
        doc,
        f"M2 had lower anomaly MAE and higher Spearman rank correlation on all "
        f"{len(metrics['dates'])} usable final-test dates.",
        bullet_num_id,
    )
    add_list_item(
        doc,
        "The correct conclusion is promising held-out predictive signal without "
        "full protocol-level confirmation.",
        bullet_num_id,
    )

    doc.add_page_break()
    doc.add_heading("1. Introduction", level=1)
    add_body_paragraph(
        doc,
        "Heat within a city is not spatially uniform. Surface cover, urban form, "
        "topography, distance from the coast, vegetation, water, and weather can "
        "produce substantial neighborhood-scale contrasts. The Local Climate "
        "Zone framework formalizes the relevance of surface structure and cover "
        "to urban thermal environments [10], while satellite studies have shown "
        "associations between vegetation abundance and land-surface temperature "
        "(LST) [11]. These relationships motivate prediction, but they do not by "
        "themselves establish causal effects for Los Angeles."
    )
    add_body_paragraph(
        doc,
        "Landsat thermal observations provide spatially detailed LST, yet any "
        "single overpass is constrained by revisit timing, clouds, atmospheric "
        "conditions, and local quality screening. A model built from routinely "
        "available public inputs could help estimate the surface-heat pattern on "
        "dates when a thermal image is not supplied to the predictor. Such a "
        "model might also help rank neighborhoods by relative surface-heat "
        "hazard. The scientific target must remain precise: satellite LST is a "
        "surface radiometric quantity and does not equal the air temperature or "
        "human exposure experienced at a particular place and time [1, 9]."
    )
    add_body_paragraph(
        doc,
        "Urban heat prediction is especially vulnerable to optimistic validation. "
        "Nearby tracts and neighboring dates share environmental structure, so a "
        "random split of tract-date rows would place strongly related examples in "
        "both training and test sets. Structured spatial and temporal validation "
        "is therefore required [7]. In addition, predictors must be available at "
        "the stated prediction origin. Same-day thermal values, same-scene "
        "optical data, future weather, and target-derived statistics would leak "
        "information."
    )
    add_body_paragraph(
        doc,
        "This study asked: can public weather, land-use and geographic, calendar, "
        "and lagged nonthermal satellite features predict neighborhood-level "
        "daytime surface-heat risk in Los Angeles? We evaluated this as a "
        "historical hindcast, first with grouped development data from 2020-2024 "
        "and then with a single prespecified, frozen 2025 test. The primary "
        "comparison was a legal ridge baseline (B1) against a nonlinear "
        "all-feature model (M2)."
    )

    doc.add_heading("2. Methods", level=1)
    doc.add_heading("2.1 Study setting and unit of analysis", level=2)
    add_body_paragraph(
        doc,
        "The study domain was the City of Los Angeles. The frozen spatial "
        "universe contained 1,096 census tracts under the project's boundary and "
        "special-use rules. The unit of analysis was a census tract crossed with "
        "a physical Landsat overpass date. Adjacent Landsat scenes from the same "
        "physical overpass were mosaicked and counted as one date, preventing "
        "duplicate temporal units. The study season was May through October. "
        "Development covered 2020-2024; calendar year 2025 was reserved for the "
        "one-time final test."
    )

    doc.add_heading("2.2 Outcome: QA-filtered daytime Landsat LST", level=2)
    add_body_paragraph(
        doc,
        "The outcome was median daytime LST for each tract-date, derived from "
        "USGS Landsat 8-9 Collection 2 Level-2 Surface Temperature products [1]. "
        "Digital numbers were converted using the documented scale and offset "
        "(temperature in kelvin = DN x 0.00341802 + 149), followed by conversion "
        "to degrees Celsius. Pixel QA excluded fill, dilated cloud, cirrus, "
        "cloud, cloud shadow, snow, and water flags, saturated pixels, terrain "
        "occlusion, and pixels closer than 1 km to clouds. No global scene-cloud "
        "cutoff determined label availability."
    )
    add_body_paragraph(
        doc,
        "A WorldCover-derived eligible-land denominator was fixed for every tract "
        "and held invariant across dates. A tract-date required at least 20 valid "
        "pixels, at least 60% valid eligible-land coverage, and at least 90% scene "
        "footprint coverage. A date required at least 98% union city footprint "
        "coverage and at least 50% tract retention. The 2 K ST_QA threshold was "
        "reserved for a sensitivity analysis rather than imposed on the primary "
        "target."
    )

    doc.add_heading("2.3 Predictors and timing contract", level=2)
    add_body_paragraph(
        doc,
        "The prediction origin was 00:00 local time on the target date. Dynamic "
        "observed inputs ended on target day -1. The primary analysis therefore "
        "represents a historical hindcast, not a live operational weather "
        "forecast. Landsat thermal bands, LST products, target-derived fields, "
        "same-scene optical data, future observations, and tract identifiers were "
        "prohibited as predictors."
    )
    add_table_caption(doc, "Table 1. Frozen M2 predictor families and timing")
    add_table(
        doc,
        headers=("Feature family", "Count", "Examples", "Availability contract"),
        rows=(
            (
                "Static land-use / geography",
                "18",
                "NLCD cover fractions; imperviousness; elevation; slope; coast distance",
                "Fixed tract summaries; raw coordinates excluded",
            ),
            (
                "Calendar",
                "2",
                "Sine and cosine of day of year",
                "Known at prediction origin",
            ),
            (
                "Lagged Daymet weather",
                "21",
                "Day length, precipitation, radiation, Tmax, Tmin, vapor pressure, solar energy",
                "1-, 3-, and 7-day windows, each ending on day -1",
            ),
            (
                "Lagged Sentinel-2 optical",
                "5",
                "NDVI, EVI, NDWI, NDBI, broadband albedo proxy",
                "Median composite from day -60 through day -1",
            ),
        ),
        widths_dxa=(1900, 800, 3280, 3380),
        numeric_columns={1},
    )
    add_source_note(
        doc,
        "Sentinel-2 L2A processing and the five optical indices follow the "
        "documented project sources [2-6]. Daymet weather is the 1 km daily V4 "
        "R1 product [8].",
    )

    doc.add_heading("2.4 Models", level=2)
    add_body_paragraph(
        doc,
        "B1 was a ridge regression baseline using 23 features: two calendar "
        "terms and 21 lagged Daymet weather terms. Its frozen regularization "
        "parameter was alpha = 10. Features were standardized within each "
        "training fold. M2 was a histogram-gradient-boosting regressor with "
        "absolute-error loss and all 46 predictors. Its frozen parameters were "
        "learning rate 0.05, 300 boosting iterations, 31 maximum leaf nodes, "
        "minimum 50 samples per leaf, and L2 regularization 1.0. Dynamic-feature "
        "median imputation was fit within training data only. Tract identifiers "
        "were never predictors, and the primary model did not use raw "
        "coordinates."
    )
    add_body_paragraph(
        doc,
        "B1 and M2 differ in both learning algorithm and feature set. Their "
        "performance contrast therefore tests a prespecified predictive package, "
        "not the causal contribution of any single land-use, weather, or optical "
        "feature."
    )

    doc.add_heading("2.5 Grouped validation and model freezing", level=2)
    add_body_paragraph(
        doc,
        "Development validation held out whole years, fixed contiguous 5 km "
        "spatial blocks, and joint year-by-block combinations. Preprocessing, "
        "imputation, scaling, and model fitting occurred inside the training fold. "
        "The primary development comparison used the joint split. After "
        "development gates were met, the selected B1 and M2 configurations were "
        "refit on all legal 2020-2024 rows and frozen before any 2025 target value "
        "was opened."
    )
    add_body_paragraph(
        doc,
        "The one-time 2025 procedure first authenticated the model lock and "
        "25,208 target-blind predictor rows across 23 physical overpasses, then "
        "froze predictions, recorded the value-opening boundary, built all 23 "
        "target caches, and atomically published the exact evaluation output set. "
        "No model, threshold, or feature rule was changed after 2025 values were "
        "opened."
    )

    doc.add_heading("2.6 Metrics and uncertainty", level=2)
    add_body_paragraph(
        doc,
        "The primary metric was equal-date-weighted mean absolute error (MAE), "
        "which prevents dates with more retained tracts from dominating. "
        "Secondary metrics were pooled RMSE, pooled out-of-sample R2, "
        "equal-date-weighted within-date anomaly MAE, and median per-date "
        "Spearman rank correlation. The exact top-20% hotspot endpoint was "
        "evaluated only on dates passing the frozen spatial-representativeness "
        "gate; ties were broken by ascending tract GEOID."
    )
    add_body_paragraph(
        doc,
        "Uncertainty used 5,000 paired crossed date-by-spatial-block bootstrap "
        "replicates. Complete dates and complete spatial blocks were sampled, "
        "never individual tract-date rows. The frozen final success rule required "
        "at least 10% point MAE improvement, median per-date Spearman at least "
        "0.50, and a 95% relative-improvement interval with lower bound greater "
        "than zero."
    )

    doc.add_page_break()
    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Development evaluation, 2020-2024", level=2)
    add_body_paragraph(
        doc,
        f"The development cohort contained {dev['tract_date_row_count']:,} "
        f"tract-date rows, {dev['independent_date_count']} independent dates, and "
        f"{dev['independent_spatial_block_count']} spatial blocks. B1 achieved "
        f"equal-date MAE {dev['baseline_point_mae_c']:.4f} C; M2 achieved "
        f"{dev['target_model_point_mae_c']:.4f} C. The improvement was "
        f"{dev['relative_mae_improvement_percent']:.2f}% "
        f"(95% crossed-cluster CI "
        f"{dev['relative_mae_improvement_ci_lower_percent']:.2f}% to "
        f"{dev['relative_mae_improvement_ci_upper_percent']:.2f}%; "
        f"P[improvement > 0] = {dev['probability_improvement_gt_zero']:.3f}). "
        f"Median per-date Spearman for M2 was "
        f"{dev_gates['observed_median_per_date_spearman']:.4f}. The development "
        "protocol gates passed, although the stronger claim that the full interval "
        "exceeded 10% did not."
    )
    add_figure(
        doc,
        FIGURES / "model_diagnostics" / "joint_performance_overview.png",
        width_inches=6.35,
        alt_text=(
            "Three-panel development performance figure comparing B1 and M2 "
            "for joint MAE, hotspot metrics, and Landsat sensor strata."
        ),
        caption=(
            "Figure 1. Grouped development performance. Panel A shows the "
            "joint spatiotemporal MAE comparison and crossed-cluster interval. "
            "Panel B shows the coverage-gated top-20% hotspot endpoint. Panel C "
            "shows sensor-stratified development error. All values are "
            "out-of-fold predictions from 2020-2024."
        ),
        source_note=(
            "Source: authenticated development tables and "
            "reports/figures/generated/model_diagnostics/"
            "joint_performance_overview.png."
        ),
    )
    gated_date_count = metrics["development_endpoints"]["relative_endpoint"][
        "gated_independent_date_count"
    ]
    add_body_paragraph(
        doc,
        f"On {gated_date_count} "
        "coverage-gated development dates, mean per-date average precision rose "
        f"from {dev_hot['B1']['mean_per_date_average_precision']:.3f} to "
        f"{dev_hot['M2']['mean_per_date_average_precision']:.3f}, and exact "
        f"top-20% recall rose from "
        f"{dev_hot['B1']['mean_per_date_recall_at_k']:.3f} to "
        f"{dev_hot['M2']['mean_per_date_recall_at_k']:.3f}. Development ablation "
        "refits also favored the full model over calendar-plus-weather, "
        "calendar-plus-static, and calendar-plus-lagged-satellite subsets, but "
        "these remain predictive associations rather than causal attributions."
    )

    doc.add_heading("3.2 Frozen one-time 2025 evaluation", level=2)
    add_body_paragraph(
        doc,
        "The frozen universe contained 25,208 tract-date keys from 1,096 tracts "
        "and 23 physical overpasses. Targets were available for 16,378 rows. "
        "After the unchanged date-level retention rule, the formal cohort "
        "contained 15,116 rows, 1,084 represented tracts, 15 independent dates, "
        "and 71 spatial blocks. Eight overpass dates were excluded from the "
        "absolute-LST evaluation because they failed the prespecified date rule."
    )
    add_table_caption(doc, "Table 2. Primary and secondary held-out metrics")
    add_table(
        doc,
        headers=("Metric", "B1", "M2", "Direction favored"),
        rows=(
            (
                "Equal-date-weighted MAE (C)",
                f"{float(b1['equal_date_weighted_mae_c']):.4f}",
                f"{float(m2['equal_date_weighted_mae_c']):.4f}",
                "Lower",
            ),
            (
                "Pooled RMSE (C)",
                f"{float(b1['pooled_rmse_c']):.4f}",
                f"{float(m2['pooled_rmse_c']):.4f}",
                "Lower",
            ),
            (
                "Pooled out-of-sample R2",
                f"{float(b1['pooled_oos_r2']):.4f}",
                f"{float(m2['pooled_oos_r2']):.4f}",
                "Higher",
            ),
            (
                "Within-date anomaly MAE (C)",
                f"{float(b1['equal_date_weighted_within_date_anomaly_mae_c']):.4f}",
                f"{float(m2['equal_date_weighted_within_date_anomaly_mae_c']):.4f}",
                "Lower",
            ),
            (
                "Median per-date Spearman",
                f"{float(b1['median_per_date_spearman']):.4f}",
                f"{float(m2['median_per_date_spearman']):.4f}",
                "Higher",
            ),
        ),
        widths_dxa=(3900, 1500, 1500, 2460),
        numeric_columns={1, 2, 3},
    )
    add_source_note(
        doc,
        "Source: frozen final_evaluation/model_metrics.csv; 15,116 rows, "
        "15 independent dates, 71 spatial blocks.",
    )
    add_body_paragraph(
        doc,
        f"M2 reduced equal-date MAE by {boot['absolute_mae_improvement_c']:.4f} C, "
        f"or {boot['relative_mae_improvement_percent']:.2f}%. The paired "
        f"crossed-cluster 95% interval was "
        f"{boot['absolute_mae_improvement_ci_lower_c']:.4f} to "
        f"{boot['absolute_mae_improvement_ci_upper_c']:.4f} C in absolute terms "
        f"and {boot['relative_mae_improvement_ci_lower_percent']:.2f}% to "
        f"{boot['relative_mae_improvement_ci_upper_percent']:.2f}% in relative "
        f"terms. The bootstrap fraction above zero improvement was "
        f"{boot['probability_improvement_gt_zero']:.4f}."
    )
    add_callout(
        doc,
        "Primary result",
        "M2 showed a large favorable point estimate, but the 95% interval still "
        "included no improvement. The prespecified final protocol therefore did "
        "not declare success.",
        forest=True,
    )
    add_table_caption(doc, "Table 3. Frozen final-test success gates")
    gate_rows = []
    for row in metrics["gates"]:
        if row["required_for_protocol_success"] != "True":
            continue
        label = {
            "median_per_date_spearman": "Median per-date Spearman",
            "point_relative_mae_improvement": "Point relative MAE improvement",
            "uncertainty_supports_positive_improvement": "Relative CI lower bound",
        }[row["gate_id"]]
        observed = float(row["observed_value"])
        threshold = float(row["threshold"])
        is_percent = row["gate_id"] != "median_per_date_spearman"
        gate_rows.append(
            (
                label,
                f"{observed * 100:.2f}%" if is_percent else f"{observed:.4f}",
                f"{row['comparison']} "
                + (f"{threshold * 100:.0f}%" if is_percent else f"{threshold:.2f}"),
                "Pass" if row["passed"] == "True" else "Fail",
            )
        )
    add_table(
        doc,
        headers=("Gate", "Observed", "Requirement", "Result"),
        rows=gate_rows,
        widths_dxa=(3900, 1700, 2200, 1560),
        numeric_columns={1, 2, 3},
    )
    add_source_note(
        doc,
        "Source: frozen final_evaluation/protocol_gates.csv. Overall protocol "
        "success = false.",
    )

    doc.add_heading("3.3 Date-to-date heterogeneity", level=2)
    add_body_paragraph(
        doc,
        f"M2 had lower absolute MAE on {metrics['lower_abs_dates']} of "
        f"{len(metrics['dates'])} dates, lower within-date anomaly MAE on "
        f"{metrics['lower_anomaly_dates']} of {len(metrics['dates'])}, and higher "
        f"Spearman rank correlation on {metrics['higher_spearman_dates']} of "
        f"{len(metrics['dates'])}. Its absolute MAE was worse on three dates. "
        "The strongest aggregate driver was 21 October 2025, when B1 MAE was "
        "12.0512 C and M2 MAE was 1.2999 C. Complete-date resampling repeatedly "
        "included or omitted such dates, producing the wide uncertainty interval."
    )
    add_figure(
        doc,
        FINAL / "per_date_error_and_rank.png",
        width_inches=6.35,
        alt_text=(
            "Per-date lines show B1 and M2 MAE and Spearman rank correlation "
            "across fifteen usable 2025 overpass dates."
        ),
        caption=(
            "Figure 2. Held-out performance by physical overpass date. Solid "
            "lines use the left axis for MAE; dashed lines use the right axis "
            "for Spearman rank correlation. The large B1 error on 21 October "
            "shows why date-level heterogeneity matters."
        ),
        source_note=(
            "Source: frozen final_evaluation/per_date_metrics.csv and "
            "per_date_error_and_rank.png."
        ),
    )

    doc.add_page_break()
    doc.add_heading("3.4 Spatial pattern and residuals", level=2)
    add_body_paragraph(
        doc,
        "Across matched available dates per tract, M2 reproduced the broad "
        "north-south and inland-coastal surface-temperature pattern more closely "
        "than B1. The B1 residual map contains broad coherent regions of "
        "overprediction and underprediction; M2 residuals are visually smaller "
        "and less spatially structured, although nonzero local errors remain. "
        "These maps average over each tract's available frozen dates and must not "
        "be read as a single-day exposure map."
    )
    add_figure(
        doc,
        map_composite,
        width_inches=6.25,
        alt_text=(
            "Four Los Angeles tract maps show observed LST, M2 predicted LST, "
            "B1 residual, and M2 residual across matched usable 2025 dates."
        ),
        caption=(
            "Figure 3. Frozen held-out spatial summaries. Top left: observed "
            "QA-filtered LST. Top right: M2 prediction. Bottom left: B1 "
            "residual (prediction minus observed). Bottom right: M2 residual. "
            "Observed and predicted panels share a temperature scale; residual "
            "panels share a diverging scale."
        ),
        source_note=(
            "Source: authenticated tract geometry and frozen "
            "final_evaluation/tract_choropleth_summary.csv; panels rendered from "
            "observed_predicted_residual_maps.pdf."
        ),
    )

    doc.add_heading("3.5 Neighborhood hotspot ranking", level=2)
    hot_b1 = metrics["hotspot"]["B1"]
    hot_m2 = metrics["hotspot"]["M2"]
    add_body_paragraph(
        doc,
        "The relative hottest-20% endpoint passed its spatial coverage gate on "
        f"{hot_m2['independent_date_count']} dates, representing "
        f"{int(hot_m2['tract_date_row_count']):,} tract-date rows. Mean "
        f"per-date average precision increased from "
        f"{float(hot_b1['mean_per_date_average_precision']):.4f} for B1 to "
        f"{float(hot_m2['mean_per_date_average_precision']):.4f} for M2. "
        f"Exact-k precision and recall each increased from "
        f"{float(hot_b1['mean_per_date_precision_at_k']):.4f} to "
        f"{float(hot_m2['mean_per_date_precision_at_k']):.4f}. These ranking "
        "results support potential screening utility, but they do not override "
        "the failed primary uncertainty gate."
    )
    add_figure(
        doc,
        FINAL / "hotspot_precision_recall.png",
        width_inches=6.35,
        alt_text=(
            "Per-date hotspot diagnostics show B1 and M2 average precision, "
            "precision at k, and recall at k across ten gated 2025 dates."
        ),
        caption=(
            "Figure 4. Frozen exact-top-20% neighborhood hotspot diagnostics. "
            "Continuous prediction scores determine average precision; exact-k "
            "selection uses descending score with ascending GEOID as the frozen "
            "tie-break."
        ),
        source_note=(
            "Source: frozen final_evaluation/hotspot_per_date.csv, "
            "hotspot_summary.csv, and hotspot_precision_recall.png."
        ),
    )

    doc.add_heading("3.6 Sensor and missingness checks", level=2)
    sensor_lookup = {
        (row["model_id"], row["sensor"]): row for row in metrics["sensors"]
    }
    sent_lookup = {
        (row["model_id"], row["sentinel_stratum"]): row
        for row in metrics["sentinel"]
    }
    add_table_caption(doc, "Table 4. Prespecified held-out subgroup checks")
    add_table(
        doc,
        headers=("Stratum", "Dates", "Rows", "B1 MAE (C)", "M2 MAE (C)"),
        rows=(
            (
                "Landsat 8",
                sensor_lookup[("M2", "landsat-8")]["independent_date_count"],
                f"{int(sensor_lookup[('M2', 'landsat-8')]['tract_date_row_count']):,}",
                f"{float(sensor_lookup[('B1', 'landsat-8')]['equal_date_weighted_mae_c']):.4f}",
                f"{float(sensor_lookup[('M2', 'landsat-8')]['equal_date_weighted_mae_c']):.4f}",
            ),
            (
                "Landsat 9",
                sensor_lookup[("M2", "landsat-9")]["independent_date_count"],
                f"{int(sensor_lookup[('M2', 'landsat-9')]['tract_date_row_count']):,}",
                f"{float(sensor_lookup[('B1', 'landsat-9')]['equal_date_weighted_mae_c']):.4f}",
                f"{float(sensor_lookup[('M2', 'landsat-9')]['equal_date_weighted_mae_c']):.4f}",
            ),
            (
                "Sentinel complete",
                sent_lookup[("M2", "sentinel_complete")]["independent_date_count"],
                f"{int(sent_lookup[('M2', 'sentinel_complete')]['tract_date_row_count']):,}",
                f"{float(sent_lookup[('B1', 'sentinel_complete')][
                    'equal_date_weighted_mae_c'
                ]):.4f}",
                f"{float(sent_lookup[('M2', 'sentinel_complete')][
                    'equal_date_weighted_mae_c'
                ]):.4f}",
            ),
            (
                "All five Sentinel features missing",
                sent_lookup[("M2", "sentinel_all_five_missing")][
                    "independent_date_count"
                ],
                f"{int(sent_lookup[('M2', 'sentinel_all_five_missing')][
                    'tract_date_row_count'
                ]):,}",
                f"{float(sent_lookup[('B1', 'sentinel_all_five_missing')][
                    'equal_date_weighted_mae_c'
                ]):.4f}",
                f"{float(sent_lookup[('M2', 'sentinel_all_five_missing')][
                    'equal_date_weighted_mae_c'
                ]):.4f}",
            ),
        ),
        widths_dxa=(3280, 1000, 1400, 1840, 1840),
        numeric_columns={1, 2, 3, 4},
    )
    add_source_note(
        doc,
        "Source: frozen final_evaluation/sensor_summary.csv and "
        "sentinel_stratum_summary.csv. Subgroups are descriptive and have no "
        "multiplicity-adjusted inferential claim.",
    )
    add_body_paragraph(
        doc,
        "Performance was not uniform by sensor. Landsat 8 dates favored M2 "
        "(3.5822 C to 1.9768 C), whereas Landsat 9 MAE changed slightly in the "
        "wrong direction (2.4179 C to 2.4471 C), even though Landsat 9 rank "
        "correlation favored M2. The all-five-Sentinel-missing stratum contained "
        "only 273 rows on two dates and 29 blocks, too little support for a "
        "general conclusion."
    )

    doc.add_heading("3.7 Target-QA sensitivity", level=2)
    add_body_paragraph(
        doc,
        "The prespecified pixel-level ST_QA <= 2 K sensitivity rebuild completed "
        "all 90 development dates, but only 15 passed the unchanged usable-date "
        "rule, below the required 30. On those 15 retained dates, the frozen "
        "primary out-of-fold predictions showed a 17.8% M2 improvement, but the "
        "95% crossed-cluster interval (-5.2% to 39.6%) crossed zero. The strict "
        "target was therefore not promoted."
    )
    add_figure(
        doc,
        FIGURES / "model_diagnostics" / "qa_cohort_improvement_forest.png",
        width_inches=6.10,
        alt_text=(
            "Forest plot compares M2 relative MAE improvement across the primary "
            "development cohort and QA sensitivity cohorts."
        ),
        caption=(
            "Figure 5. Development target-QA sensitivity. Intervals summarize "
            "crossed date-by-spatial-block uncertainty. The strict pixel-level "
            "ST_QA cohort retained too few independent dates and its interval "
            "crossed zero."
        ),
        source_note=(
            "Source: authenticated development QA diagnostics and "
            "reports/figures/generated/model_diagnostics/"
            "qa_cohort_improvement_forest.png."
        ),
    )

    doc.add_page_break()
    doc.add_heading("4. Discussion", level=1)
    doc.add_heading("4.1 What the study supports", level=2)
    add_body_paragraph(
        doc,
        "The frozen 2025 point estimate indicates that public calendar, weather, "
        "land-use, geographic, and lagged optical information can carry useful "
        "signal for neighborhood-scale surface LST in Los Angeles. The result is "
        "not limited to pooled error: M2 also improved within-date anomaly "
        "accuracy, neighborhood rank correlation, and top-20% hotspot retrieval. "
        "The spatial maps show that the nonlinear all-feature model captured broad "
        "temperature structure while reducing large coherent baseline residuals."
    )
    add_body_paragraph(
        doc,
        "The held-out estimate is directionally consistent with the development "
        "result. Development produced a smaller but statistically supported "
        "16.19% improvement, while the one-time 2025 point estimate was 30.53%. "
        "This consistency matters, but it does not erase the final interval. The "
        "prespecified inferential unit included independent dates and spatial "
        "blocks, so 15 usable dates provide much less information than 15,116 "
        "tract-date rows might suggest."
    )
    doc.add_heading("4.2 Why the final gate failed", level=2)
    add_body_paragraph(
        doc,
        "The failed uncertainty gate reflects temporal heterogeneity rather than "
        "a contradiction in the point estimates. M2 improved absolute MAE on 12 "
        "dates, but three dates moved in the opposite direction, and the unusually "
        "large B1 error on 21 October contributed strongly to the average "
        "improvement. When complete dates are resampled, the resulting estimate "
        "varies substantially. This is exactly the dependence structure the "
        "prespecified crossed bootstrap was designed to preserve."
    )
    add_callout(
        doc,
        "Claim discipline",
        "The study may report the 30.53% held-out point reduction and the stronger "
        "ranking metrics, but it must also report the -10.13% to 58.46% interval "
        "and the failed overall protocol flag.",
    )
    doc.add_heading("4.3 Practical meaning", level=2)
    add_body_paragraph(
        doc,
        "A model with this pattern of performance may be useful for research "
        "screening: estimating relative surface-heat patterns, prioritizing areas "
        "for follow-up measurement, or testing where additional sensors would be "
        "most informative. It is not yet an operational product. The present "
        "weather inputs are lagged historical observations, not a live forecast "
        "feed, and the 2025 evaluation includes only dates that supplied a usable "
        "Landsat target for scoring. Deployment would require prospective data "
        "latency tests, forecast-compatible meteorology, monitoring for sensor "
        "shift, and evaluation on future dates."
    )
    doc.add_heading("4.4 Scientific interpretation", level=2)
    add_body_paragraph(
        doc,
        "The model predicts a surface-heat hazard proxy. Even a highly accurate "
        "LST model cannot directly determine the air temperature a resident "
        "breathes, indoor conditions, time-activity patterns, physiologic "
        "vulnerability, heat illness, or mortality. Satellite LST can help "
        "describe spatial surface conditions, but its relationship with short-"
        "timescale human exposure is incomplete [9]. Any public-health use would "
        "require separate exposure and outcome validation."
    )

    doc.add_heading("5. Limitations", level=1)
    limitations = (
        "Outcome boundary. Daytime clear-sky Landsat LST is not air temperature, "
        "personal exposure, or a health outcome.",
        "Selective observability. Clouds, scene footprint, pixel QA, and tract "
        "retention reduced 23 frozen overpasses to 15 usable evaluation dates.",
        "Limited independent time support. The primary interval is wide because "
        "the effective temporal sample is 15 dates, not 15,116 independent rows.",
        "Date and sensor heterogeneity. M2 was worse on three dates, and Landsat "
        "9 absolute MAE did not improve.",
        "Confounded model comparison. B1 and M2 differ in both algorithm and "
        "feature set, so the contrast cannot isolate one predictor family's effect.",
        "No causal interpretation. Land cover, vegetation, imperviousness, and "
        "other features are predictive associations, not intervention effects.",
        "Spatial aggregation. Census-tract medians hide within-tract variation and "
        "depend on the modifiable areal unit.",
        "Geographic and temporal scope. The frozen final test covers one city and "
        "one calendar year; transportability to other climates or future years is "
        "unknown.",
        "Residual dependence. Development residuals remained spatially clustered, "
        "even though M2 reduced mean Moran's I relative to B1.",
        "Operational gap. The design is a historical hindcast using observations "
        "available through day -1, not a prospective forecast system.",
    )
    for item in limitations:
        add_list_item(doc, item, bullet_num_id)

    doc.add_heading("6. Data, Evidence, and Reproducibility", level=1)
    add_body_paragraph(
        doc,
        "All reported values are generated from frozen authenticated artifacts; no "
        "table or figure value in this paper was hand-edited. The model lock binds "
        "the exact B1 and M2 pipelines, feature order, training keys, grouped "
        "validation design, hotspot rule, and figure plan. The final transaction "
        "records readiness, authorization, one consumption claim, frozen "
        "predictions, the values-opened boundary, 23 target-cache commits, an exact "
        "21-file result directory, and a completion marker."
    )
    add_body_paragraph(
        doc,
        "The read-only evidence export contains 239 files. Its ZIP SHA-256 is "
        "61a853c3eeea3f1ae92bf7999f0fd057018797f70498fcd017d1394dbd621b51. "
        "The verifier authenticated the six-state transaction chain, all 23 cache "
        "commits, the exact final output set, recovery records, and an independently "
        "checked repository bundle. Re-running the original evaluation command "
        "after publication entered completion-authentication mode and returned the "
        "same completion commit without recomputing the evaluation."
    )
    add_body_paragraph(
        doc,
        "A read-only interactive visualization of observed LST, M2 and B1 "
        "predictions, residuals, per-date diagnostics, and uncertainty is available "
        "at https://cmschase.github.io/LA-neighborhood-heat/. The public GitHub "
        "Pages site is a display layer over the frozen outputs, not an independent "
        "model run."
    )
    add_body_paragraph(
        doc,
        "Reproduction should begin from the canonical repository manifests and "
        "configuration, not from screenshots or manually copied values. The "
        "frozen final evaluation must not be retuned or repeated as a new claim. "
        "Future extensions should register a new prospective protocol for "
        "additional years or cities while preserving the 2025 result unchanged."
    )

    doc.add_heading("7. Conclusion", level=1)
    add_body_paragraph(
        doc,
        "Public weather, land-use and geographic, calendar, and lagged optical "
        "satellite features showed useful held-out ability to predict "
        "neighborhood-scale daytime Landsat LST in Los Angeles. In the frozen "
        "2025 test, M2 reduced point-estimate MAE from 3.1165 C to 2.1650 C and "
        "substantially improved neighborhood ranking and hotspot retrieval. "
        "However, the crossed date-by-block 95% interval for relative MAE "
        "improvement extended from -10.13% to 58.46%, so the prespecified "
        "positive-lower-bound gate failed. The scientifically correct conclusion "
        "is therefore promising held-out predictive signal without protocol-level "
        "confirmation. Stronger evidence will require more independent dates, "
        "future-year prospective testing, external-city validation, and separate "
        "evaluation against air-temperature or human-exposure measurements if "
        "those are the intended applications."
    )

    doc.add_page_break()
    doc.add_heading("References", level=1)
    references = (
        "U.S. Geological Survey. Landsat 8-9 Collection 2 Level-2 Science Product "
        "Guide and Landsat Collection 2 Surface Temperature documentation. Data "
        "DOI: https://doi.org/10.5066/P9OGBGM6.",
        "European Space Agency / Copernicus. Sentinel-2 Level-2A processing "
        "documentation. Data DOI: https://doi.org/10.5270/S2_-znk9xsj.",
        "Huete, A., Didan, K., Miura, T., Rodriguez, E. P., Gao, X., & Ferreira, "
        "L. G. (2002). Overview of the radiometric and biophysical performance "
        "of the MODIS vegetation indices. Remote Sensing of Environment, 83, "
        "195-213. https://doi.org/10.1016/S0034-4257(02)00096-2.",
        "McFeeters, S. K. (1996). The use of the Normalized Difference Water "
        "Index (NDWI) in the delineation of open water features. International "
        "Journal of Remote Sensing, 17, 1425-1432. "
        "https://doi.org/10.1080/01431169608948714.",
        "Zha, Y., Gao, J., & Ni, S. (2003). Use of normalized difference built-up "
        "index in automatically mapping urban areas from TM imagery. "
        "International Journal of Remote Sensing, 24, 583-594. "
        "https://doi.org/10.1080/01431160304987.",
        "Bonafoni, S., & Sekertekin, A. (2020). Albedo retrieval from Sentinel-2 "
        "by new narrow-to-broadband conversion coefficients. IEEE Geoscience and "
        "Remote Sensing Letters. https://doi.org/10.1109/LGRS.2020.2967085.",
        "Roberts, D. R., Bahn, V., Ciuti, S., Boyce, M. S., Elith, J., Guillera-"
        "Arroita, G., et al. (2017). Cross-validation strategies for data with "
        "temporal, spatial, hierarchical, or phylogenetic structure. Ecography, "
        "40, 913-929. https://doi.org/10.1111/ecog.02881.",
        "Thornton, M. M., Shrestha, R., Wei, Y., Thornton, P. E., Kao, S.-C., & "
        "Wilson, B. E. Daymet: Daily Surface Weather Data on a 1-km Grid for "
        "North America, Version 4 R1. ORNL DAAC. "
        "https://doi.org/10.3334/ORNLDAAC/2129.",
        "White-Newsome, J. L., Brines, S. J., Brown, D. G., Dvonch, J. T., "
        "Gronlund, C. J., Zhang, K., et al. (2013). Validating satellite-derived "
        "land surface temperature with in situ measurements: a public health "
        "perspective. Environmental Health Perspectives, 121, 925-931. "
        "https://doi.org/10.1289/ehp.1206176.",
        "Stewart, I. D., & Oke, T. R. (2012). Local Climate Zones for urban "
        "temperature studies. Bulletin of the American Meteorological Society, "
        "93, 1879-1900. https://doi.org/10.1175/BAMS-D-11-00019.1.",
        "Weng, Q., Lu, D., & Schubring, J. (2004). Estimation of land surface "
        "temperature-vegetation abundance relationship for urban heat island "
        "studies. Remote Sensing of Environment, 89, 467-483. "
        "https://doi.org/10.1016/j.rse.2003.11.005.",
        "Moran, P. A. P. (1950). Notes on continuous stochastic phenomena. "
        "Biometrika, 37, 17-23. https://doi.org/10.1093/biomet/37.1-2.17.",
    )
    for index, text in enumerate(references, start=1):
        paragraph = doc.add_paragraph(style="Reference")
        paragraph.add_run(f"{index}. {text}")

    doc.add_heading("Appendix A. Frozen evaluation identity", level=1)
    identity_rows = (
        (
            "Consumption claim",
            "c174e0b26272dcb194a54ec4cdb468e18d0f64f8d04156681746a52361d1f01f",
        ),
        (
            "Completion commit",
            "4cc8a5536cf1055d42876577f8d9f6300c799176779a7ec89cd1d3ed819d77a0",
        ),
        (
            "Model lock file SHA-256",
            "bf77762bbd1838be2b67e8461c5f99aad1c2ebf36b4f3b53b25dac1801a81245",
        ),
        (
            "Evidence ZIP SHA-256",
            "61a853c3eeea3f1ae92bf7999f0fd057018797f70498fcd017d1394dbd621b51",
        ),
        ("Final cohort", "15,116 tract-date rows; 15 dates; 71 spatial blocks"),
        ("Frozen protocol success", "False"),
    )
    add_table_caption(doc, "Table A1. Immutable identifiers and audit anchors")
    add_table(
        doc,
        headers=("Field", "Value"),
        rows=identity_rows,
        widths_dxa=(2600, 6760),
    )

    core = doc.core_properties
    core.title = (
        "Can Public Weather, Land-Use, and Lagged Optical Satellite Features "
        "Predict Neighborhood-Scale Surface Heat?"
    )
    core.subject = "Prespecified historical hindcast of Los Angeles surface LST"
    core.author = "Los Angeles Surface Heat Research Project"
    core.keywords = (
        "land-surface temperature; urban heat; Los Angeles; Landsat; "
        "Sentinel-2; Daymet; historical hindcast"
    )
    core.comments = (
        "Built from frozen authenticated project results; no final evaluation "
        "was rerun."
    )

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def _register_pdf_fonts() -> None:
    font_dir = Path("C:/Windows/Fonts")
    font_files = {
        "Calibri": font_dir / "calibri.ttf",
        "Calibri-Bold": font_dir / "calibrib.ttf",
        "Calibri-Italic": font_dir / "calibrii.ttf",
        "Calibri-BoldItalic": font_dir / "calibriz.ttf",
    }
    for name, path in font_files.items():
        if not path.exists():
            raise FileNotFoundError(path)
        if name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    pdfmetrics.registerFontFamily(
        "Calibri",
        normal="Calibri",
        bold="Calibri-Bold",
        italic="Calibri-Italic",
        boldItalic="Calibri-BoldItalic",
    )


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Normal": ParagraphStyle(
            "PaperBody",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=11,
            leading=14.66,
            alignment=TA_JUSTIFY,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=8,
            allowWidows=1,
            allowOrphans=1,
        ),
        "NormalCenter": ParagraphStyle(
            "PaperBodyCenter",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=11,
            leading=14.66,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=8,
        ),
        "NormalLeft": ParagraphStyle(
            "PaperBodyLeft",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=11,
            leading=14.66,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=8,
        ),
        "Title": ParagraphStyle(
            "PaperTitle",
            parent=base["Title"],
            fontName="Calibri-Bold",
            fontSize=29,
            leading=33,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=9,
            keepWithNext=False,
        ),
        "Subtitle": ParagraphStyle(
            "PaperSubtitle",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=14,
            leading=17,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#193C36"),
            spaceBefore=0,
            spaceAfter=8,
            keepWithNext=False,
        ),
        "CoverResult": ParagraphStyle(
            "PaperCoverResult",
            parent=base["BodyText"],
            fontName="Calibri-Bold",
            fontSize=25,
            leading=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#193C36"),
            spaceBefore=0,
            spaceAfter=14,
            keepTogether=False,
        ),
        "CoverQualifier": ParagraphStyle(
            "PaperCoverQualifier",
            parent=base["BodyText"],
            fontName="Calibri-Bold",
            fontSize=11.5,
            leading=15,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#C9583E"),
            spaceBefore=8,
            spaceAfter=40,
            keepTogether=False,
        ),
        "Heading 1": ParagraphStyle(
            "PaperH1",
            parent=base["Heading1"],
            fontName="Calibri-Bold",
            fontSize=16,
            leading=19,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=18,
            spaceAfter=10,
            keepWithNext=False,
        ),
        "Heading 2": ParagraphStyle(
            "PaperH2",
            parent=base["Heading2"],
            fontName="Calibri-Bold",
            fontSize=13,
            leading=16,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=False,
        ),
        "Heading 3": ParagraphStyle(
            "PaperH3",
            parent=base["Heading3"],
            fontName="Calibri-Bold",
            fontSize=12,
            leading=14,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=False,
        ),
        "Figure Caption": ParagraphStyle(
            "PaperFigureCaption",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=9,
            leading=10.2,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#3B4341"),
            spaceBefore=4,
            spaceAfter=8,
            keepTogether=False,
        ),
        "Source Note": ParagraphStyle(
            "PaperSourceNote",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=8,
            leading=9,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#5D6563"),
            spaceBefore=0,
            spaceAfter=4,
            keepTogether=False,
        ),
        "Table Caption": ParagraphStyle(
            "PaperTableCaption",
            parent=base["BodyText"],
            fontName="Calibri-Bold",
            fontSize=9.5,
            leading=11,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#193C36"),
            spaceBefore=4,
            spaceAfter=4,
            keepWithNext=False,
        ),
        "Reference": ParagraphStyle(
            "PaperReference",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=9.2,
            leading=10.7,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#18201F"),
            leftIndent=21.6,
            firstLineIndent=-21.6,
            spaceAfter=4,
        ),
        "List": ParagraphStyle(
            "PaperList",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=11,
            leading=13.3,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#18201F"),
            leftIndent=27,
            firstLineIndent=-14,
            bulletIndent=13,
            spaceAfter=4,
        ),
        "TableBody": ParagraphStyle(
            "PaperTableBody",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=9.2,
            leading=10.4,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=0,
        ),
        "TableBodyCenter": ParagraphStyle(
            "PaperTableBodyCenter",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=9.2,
            leading=10.4,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=0,
        ),
        "TableHead": ParagraphStyle(
            "PaperTableHead",
            parent=base["BodyText"],
            fontName="Calibri-Bold",
            fontSize=9.1,
            leading=10.4,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#193C36"),
            spaceBefore=0,
            spaceAfter=0,
        ),
        "TableHeadCenter": ParagraphStyle(
            "PaperTableHeadCenter",
            parent=base["BodyText"],
            fontName="Calibri-Bold",
            fontSize=9.1,
            leading=10.4,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#193C36"),
            spaceBefore=0,
            spaceAfter=0,
        ),
        "Callout": ParagraphStyle(
            "PaperCallout",
            parent=base["BodyText"],
            fontName="Calibri",
            fontSize=10.4,
            leading=12.3,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#18201F"),
            spaceBefore=0,
            spaceAfter=0,
        ),
    }


def _iter_docx_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def _paragraph_markup(paragraph: DocxParagraph) -> str:
    chunks: list[str] = []
    for run in paragraph.runs:
        text = escape(run.text).replace("\n", "<br/>")
        if not text:
            continue
        attributes: list[str] = []
        if run.font.size is not None:
            attributes.append(f'size="{run.font.size.pt:.1f}"')
        if run.font.color.type is not None and run.font.color.rgb is not None:
            attributes.append(f'color="#{run.font.color.rgb}"')
        if run.font.name:
            attributes.append('face="Calibri"')
        if attributes:
            text = f"<font {' '.join(attributes)}>{text}</font>"
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        chunks.append(text)
    if chunks:
        return "".join(chunks)
    return escape(paragraph.text).replace("\n", "<br/>")


def _paragraph_has_page_break(paragraph: DocxParagraph) -> bool:
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def _paragraph_is_list(paragraph: DocxParagraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.numPr is not None


def _paragraph_fill(paragraph: DocxParagraph) -> str | None:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        return None
    return shading.get(qn("w:fill"))


def _paragraph_images(
    document: Document,
    paragraph: DocxParagraph,
    *,
    image_index_start: int,
) -> tuple[list[PdfImage], int]:
    images: list[PdfImage] = []
    image_index = image_index_start
    blips = paragraph._p.xpath(".//a:blip")
    extents = paragraph._p.xpath(".//wp:extent")
    for idx, blip in enumerate(blips):
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in document.part.related_parts:
            continue
        part = document.part.related_parts[rel_id]
        extension = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/gif": ".gif",
            "image/tiff": ".tif",
        }.get(getattr(part, "content_type", ""), ".img")
        image_path = TMP_DIR / f"pdf_docx_image_{image_index:02d}{extension}"
        image_path.write_bytes(part.blob)
        image_index += 1
        with Image.open(image_path) as source:
            pixel_width, pixel_height = source.size
        if idx < len(extents):
            width = int(extents[idx].get("cx")) / 914400 * inch
            height = int(extents[idx].get("cy")) / 914400 * inch
        else:
            width = min(6.35 * inch, pixel_width / 150 * inch)
            height = width * pixel_height / pixel_width
        max_width = 6.35 * inch
        max_height = 7.15 * inch
        scale = min(1.0, max_width / width, max_height / height)
        pdf_image = PdfImage(str(image_path), width=width * scale, height=height * scale)
        pdf_image.hAlign = "CENTER"
        images.append(pdf_image)
    return images, image_index


def _numericish(text: str) -> bool:
    stripped = (
        text.strip()
        .replace(",", "")
        .replace("%", "")
        .replace("C", "")
        .replace(">=", "")
        .replace(">", "")
        .replace("<=", "")
        .replace("<", "")
        .replace("=", "")
        .strip()
    )
    if stripped in {"Pass", "Fail", "True", "False"}:
        return True
    try:
        float(stripped)
        return True
    except ValueError:
        return False


def _docx_table_to_pdf(table: DocxTable, styles: dict[str, ParagraphStyle]) -> PdfTable:
    raw_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    numeric_columns: set[int] = set()
    if raw_rows:
        for col_idx in range(len(raw_rows[0])):
            body_values = [row[col_idx] for row in raw_rows[1:] if row[col_idx].strip()]
            if body_values and all(_numericish(value) for value in body_values):
                numeric_columns.add(col_idx)

    data: list[list[PdfParagraph]] = []
    for row_idx, row in enumerate(raw_rows):
        pdf_row: list[PdfParagraph] = []
        for col_idx, value in enumerate(row):
            if row_idx == 0:
                style = (
                    styles["TableHeadCenter"]
                    if col_idx in numeric_columns
                    else styles["TableHead"]
                )
                markup = f"<b>{escape(value)}</b>"
            else:
                style = (
                    styles["TableBodyCenter"]
                    if col_idx in numeric_columns
                    else styles["TableBody"]
                )
                markup = escape(value).replace("\n", "<br/>")
            pdf_row.append(PdfParagraph(markup, style))
        data.append(pdf_row)

    first_row = table.rows[0]
    widths_dxa: list[int] = []
    for cell in first_row.cells:
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        widths_dxa.append(
            int(tc_w.get(qn("w:w"))) if tc_w is not None else TABLE_WIDTH_DXA
        )
    total = sum(widths_dxa)
    column_widths = [6.5 * inch * width / total for width in widths_dxa]
    pdf_table = PdfTable(
        data,
        colWidths=column_widths,
        repeatRows=1,
        splitByRow=1,
        hAlign="LEFT",
    )
    pdf_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D7DBE2")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return pdf_table


def build_pdf_from_docx(docx_path: Path) -> Path:
    """Render a publication PDF from the actual DOCX content without Word/LO.

    The host lacks LibreOffice, and Word's PDF automation filter is not usable.
    This converter reads the saved DOCX's paragraphs, styles, numbering, tables,
    page breaks, and embedded image relationships, then lays them out with the
    same narrative-proposal token map through ReportLab.
    """

    _register_pdf_fonts()
    styles = _pdf_styles()
    document = Document(docx_path)
    blocks = list(_iter_docx_blocks(document))
    story: list[object] = []
    image_index = 1
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if isinstance(block, DocxTable):
            story.append(_docx_table_to_pdf(block, styles))
            story.append(Spacer(1, 4))
            index += 1
            continue

        paragraph = block
        if _paragraph_has_page_break(paragraph):
            story.append(PageBreak())
            index += 1
            continue

        images, image_index = _paragraph_images(
            document,
            paragraph,
            image_index_start=image_index,
        )
        if images:
            group: list[object] = [Spacer(1, 4), *images]
            consumed = 1
            for next_offset in (1, 2):
                if index + next_offset >= len(blocks):
                    break
                candidate = blocks[index + next_offset]
                if not isinstance(candidate, DocxParagraph):
                    break
                style_name = candidate.style.name if candidate.style is not None else "Normal"
                if style_name not in {"Figure Caption", "Source Note"}:
                    break
                group.append(
                    PdfParagraph(
                        _paragraph_markup(candidate),
                        styles[style_name],
                    )
                )
                consumed += 1
            story.extend(group)
            index += consumed
            continue

        text = paragraph.text
        if not text.strip():
            story.append(Spacer(1, 8))
            index += 1
            continue

        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        style = styles.get(style_name, styles["Normal"])
        if text.endswith("lower held-out point-estimate MAE for M2 vs. B1"):
            style = styles["CoverResult"]
        elif text.startswith("Promising held-out predictive signal;"):
            style = styles["CoverQualifier"]
        elif style_name == "Normal":
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                style = styles["NormalCenter"]
            elif paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                style = styles["NormalLeft"]
            elif "http" in text or "SHA-256" in text:
                style = styles["NormalLeft"]
        markup = _paragraph_markup(paragraph)
        if _paragraph_is_list(paragraph):
            flowable = PdfParagraph(f"\u2022&nbsp;&nbsp;{markup}", styles["List"])
        else:
            flowable = PdfParagraph(markup, style)

        fill = _paragraph_fill(paragraph)
        if fill and fill not in {"auto", "FFFFFF"}:
            flowable = PdfParagraph(markup, styles["Callout"])
            callout = PdfTable([[flowable]], colWidths=[6.16 * inch], hAlign="CENTER")
            accent = "#193C36" if fill == FOREST_FILL else "#C9583E"
            callout.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{fill}")),
                        ("LINEBEFORE", (0, 0), (0, -1), 3, colors.HexColor(accent)),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ]
                )
            )
            story.extend([Spacer(1, 5), callout, Spacer(1, 8)])
        else:
            story.append(flowable)
        index += 1

    pdf_doc = BaseDocTemplate(
        str(OUTPUT_PDF),
        pagesize=LETTER,
        rightMargin=1 * inch,
        leftMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        title=document.core_properties.title,
        author=document.core_properties.author,
        subject=document.core_properties.subject,
        pageCompression=1,
    )

    def draw_page(canvas, _doc) -> None:
        canvas.saveState()
        canvas.setTitle(document.core_properties.title)
        canvas.setAuthor(document.core_properties.author)
        canvas.setSubject(document.core_properties.subject)
        if canvas.getPageNumber() > 1:
            canvas.setFont("Calibri-Bold", 8.2)
            canvas.setFillColor(colors.HexColor("#5D6563"))
            canvas.drawString(
                1 * inch,
                LETTER[1] - 0.54 * inch,
                "LA NEIGHBORHOOD SURFACE HEAT  /  RESEARCH PAPER",
            )
            canvas.setFont("Calibri", 8.5)
            canvas.drawRightString(
                LETTER[0] - 1 * inch,
                0.48 * inch,
                str(canvas.getPageNumber() - 1),
            )
        canvas.restoreState()

    frame = Frame(
        1 * inch,
        1 * inch,
        LETTER[0] - 2 * inch,
        LETTER[1] - 2 * inch,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        id="paper-frame",
    )
    pdf_doc.addPageTemplates(
        [PageTemplate(id="paper-template", frames=[frame], onPage=draw_page)]
    )
    pdf_doc.build(story)
    return OUTPUT_PDF


def main() -> None:
    output = build_document()
    pdf = build_pdf_from_docx(output)
    print(output)
    print(pdf)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"build_research_paper failed: {exc}", file=sys.stderr)
        raise
